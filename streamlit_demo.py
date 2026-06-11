from __future__ import annotations

import hashlib
import io
import os
import re
import tempfile
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import numpy as np
import pandas as pd
import streamlit as st
from rank_bm25 import BM25Okapi
from sentence_transformers import SentenceTransformer, CrossEncoder
from sklearn.metrics.pairwise import cosine_similarity

try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None

try:
    from docx import Document
except Exception:
    Document = None

try:
    import torch
    from transformers import AutoModelForCausalLM, AutoTokenizer, BitsAndBytesConfig
except Exception:
    torch = None
    AutoModelForCausalLM = None
    AutoTokenizer = None
    BitsAndBytesConfig = None


# ======================================================================================
# Report-aligned best original end-to-end pipeline
# Source-aware + Article-aware retrieval + Turkish BGE reranker + improved prompt
# + base Mistral generator.
# ======================================================================================

ROOT_DIR = Path(__file__).resolve().parent
DEFAULT_DOCS_DIR = ROOT_DIR / "final_system" / "data" / "custom_documents"

EMBEDDING_MODEL_NAME = "sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"
TURKISH_BGE_RERANKER_NAME = "seroe/bge-reranker-v2-m3-turkish-triplet"
BASE_MISTRAL_MODEL_NAME = "mistralai/Mistral-7B-Instruct-v0.2"

CANDIDATE_K = 10
FINAL_CONTEXT_K = 3
ALPHA = 0.5
HYBRID_WEIGHT = 0.7
RERANK_WEIGHT = 0.3
ARTICLE_BONUS_WEIGHT = 0.15
CHUNK_SIZE = 1200
CHUNK_OVERLAP = 200
MIN_FILTERED_ROWS = 1

SUPPORTED_UPLOAD_TYPES = ["txt", "pdf", "docx"]


# ----------------------------------
# Basic text utilities
# ----------------------------------

def normalize_whitespace(text: str) -> str:
    text = str(text).replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def simple_turkish_tokenize(text: str) -> List[str]:
    text = str(text).lower()
    text = re.sub(r"[^a-zçğıöşü0-9\s]", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text.split()


def min_max_normalize(scores: np.ndarray) -> np.ndarray:
    scores = np.asarray(scores, dtype=np.float32)
    if len(scores) == 0:
        return scores
    if float(scores.max()) == float(scores.min()):
        return np.zeros_like(scores)
    return (scores - scores.min()) / (scores.max() - scores.min())


# ----------------------------------
# Document loading
# ----------------------------------

def read_txt_from_bytes(data: bytes) -> str:
    for enc in ["utf-8", "utf-8-sig", "cp1254", "latin-1"]:
        try:
            return data.decode(enc, errors="ignore")
        except Exception:
            pass
    return data.decode("utf-8", errors="ignore")


def read_pdf_from_bytes(data: bytes) -> str:
    if fitz is None:
        raise ImportError("PDF okumak için PyMuPDF gerekli: pip install pymupdf")
    pdf = fitz.open(stream=data, filetype="pdf")
    pages = [page.get_text("text") for page in pdf]
    pdf.close()
    return "\n".join(pages)


def read_docx_from_bytes(data: bytes) -> str:
    if Document is None:
        raise ImportError("DOCX okumak için python-docx gerekli: pip install python-docx")
    doc = Document(io.BytesIO(data))
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])


def read_file_bytes(file_name: str, data: bytes) -> str:
    suffix = Path(file_name).suffix.lower()
    if suffix == ".txt":
        return read_txt_from_bytes(data)
    if suffix == ".pdf":
        return read_pdf_from_bytes(data)
    if suffix == ".docx":
        return read_docx_from_bytes(data)
    raise ValueError(f"Unsupported file type: {file_name}")


def read_file_path(path: Path) -> str:
    suffix = path.suffix.lower()
    data = path.read_bytes()
    return read_file_bytes(path.name, data)


LAW_SOURCE_RULES = [
    ("Türkiye Cumhuriyeti Anayasası", ["anayasa", "constitution", "1982 anayasası"]),
    ("Türk Ceza Kanunu", ["tck", "türk ceza kanunu", "ceza kanunu"]),
    ("Ceza Muhakemesi Kanunu", ["cmk", "ceza muhakemesi kanunu"]),
    ("Türk Medeni Kanunu", ["tmk", "medeni kanun", "türk medeni kanunu"]),
    ("Türk Borçlar Kanunu", ["tbk", "borçlar kanunu", "türk borçlar kanunu"]),
    ("Hukuk Muhakemeleri Kanunu", ["hmk", "hukuk muhakemeleri kanunu"]),
    ("İdari Yargılama Usulü Kanunu", ["iyuk", "idari yargılama usulü"]),
    ("Kişisel Verilerin Korunması Kanunu", ["kvkk", "kişisel verilerin korunması"]),
    ("Bilgi Edinme Hakkı Kanunu", ["bilgi edinme hakkı", "bilgi edinme"]),
]


def infer_source_name(file_name: str, text: str) -> str:
    haystack = f"{file_name}\n{text[:3000]}".lower()
    for source_name, keywords in LAW_SOURCE_RULES:
        if any(keyword in haystack for keyword in keywords):
            return source_name
    return Path(file_name).stem.replace("_", " ")


def detect_source_filter(query: str) -> Optional[str]:
    q = str(query).lower()

    if any(x in q for x in ["anayasa", "anayasanın", "anayasa'nın", "geçici madde", "gecici madde"]):
        return "Türkiye Cumhuriyeti Anayasası"

    # Strong explicit legal-source cues.
    source_query_rules = [
        ("Türk Ceza Kanunu", ["tck", "türk ceza kanunu", "ceza kanunu"]),
        ("Ceza Muhakemesi Kanunu", ["cmk", "ceza muhakemesi"]),
        ("Türk Medeni Kanunu", ["tmk", "medeni kanun", "miras", "gaiplik"]),
        ("Türk Borçlar Kanunu", ["tbk", "borçlar kanunu", "kira", "kiracı", "sözleşme"]),
        ("Kişisel Verilerin Korunması Kanunu", ["kvkk", "kişisel veri", "ilgili kişi"]),
        ("Bilgi Edinme Hakkı Kanunu", ["bilgi edinme"]),
        ("Hukuk Muhakemeleri Kanunu", ["hmk", "hukuk muhakemeleri"]),
        ("İdari Yargılama Usulü Kanunu", ["iyuk", "idari yargılama"]),
    ]

    for source_name, keywords in source_query_rules:
        if any(keyword in q for keyword in keywords):
            return source_name

    return None


def source_matches(source_filter: str, row: Dict[str, object]) -> bool:
    target = str(source_filter).lower()
    fields = [
        str(row.get("source", "")),
        str(row.get("file_name", "")),
        str(row.get("text", ""))[:3000],
    ]
    combined = "\n".join(fields).lower()

    # Exact-ish source name match.
    if target in combined:
        return True

    # Common aliases.
    alias_map = {
        "türkiye cumhuriyeti anayasası": ["anayasa", "constitution"],
        "türk ceza kanunu": ["tck", "ceza kanunu"],
        "ceza muhakemesi kanunu": ["cmk"],
        "türk medeni kanunu": ["tmk", "medeni kanun"],
        "türk borçlar kanunu": ["tbk", "borçlar kanunu"],
        "kişisel verilerin korunması kanunu": ["kvkk", "kişisel veri"],
        "bilgi edinme hakkı kanunu": ["bilgi edinme"],
    }
    return any(alias in combined for alias in alias_map.get(target, []))


# ----------------------------------
# Article-aware chunking and scoring
# ----------------------------------

def extract_article_reference(query: str) -> Optional[Dict[str, object]]:
    q = str(query).lower()
    is_temporary = "geçici madde" in q or "gecici madde" in q

    patterns = [
        r"geçici\s+madde\s+(\d+)",
        r"gecici\s+madde\s+(\d+)",
        r"madde\s*(\d+)",
        r"(\d+)\.\s*madde",
        r"(\d+)\s*inci\s*madde",
        r"(\d+)\s*ıncı\s*madde",
        r"(\d+)\s*uncu\s*madde",
        r"(\d+)\s*üncü\s*madde",
        r"(\d+)\.\s*maddesi",
        r"(\d+)\s*maddesi",
        r"\b(tck|cmk|tmk|tbk|hmk|kvkk)\s*(\d+)\b",
    ]

    for pattern in patterns:
        match = re.search(pattern, q)
        if match:
            number = match.group(2) if len(match.groups()) >= 2 and match.group(2) else match.group(1)
            return {"article_number": number, "is_temporary": is_temporary}

    return None


def article_match_bonus(query: str, chunk_text: str, chunk_article_number: str = "") -> float:
    ref = extract_article_reference(query)
    if ref is None:
        return 0.0

    article_no = str(ref["article_number"])
    is_temporary = bool(ref["is_temporary"])

    if chunk_article_number and str(chunk_article_number).strip() == article_no:
        return 1.0

    text = str(chunk_text).lower()
    text = re.sub(r"\[[^\]]+\]", " ", text)

    if is_temporary:
        patterns = [
            rf"geçici\s+madde\s+{article_no}\b",
            rf"gecici\s+madde\s+{article_no}\b",
            rf"geçici\s+{article_no}\.\s*madde",
            rf"gecici\s+{article_no}\.\s*madde",
            rf"geçici\s+{article_no}\s*(inci|ıncı|uncu|üncü|nci|ncı|ncu|ncü)\s+madde",
            rf"gecici\s+{article_no}\s*(inci|ıncı|uncu|üncü|nci|ncı|ncu|ncü)\s+madde",
        ]
    else:
        patterns = [
            rf"madde\s+{article_no}\b",
            rf"madde\s*{article_no}\s*[–-]",
            rf"{article_no}\.\s*madde",
            rf"{article_no}\s*maddesi",
            rf"{article_no}\s*(inci|ıncı|uncu|üncü|nci|ncı|ncu|ncü)\s+madde",
            rf"{article_no}\s*(inci|ıncı|uncu|üncü|nci|ncı|ncu|ncü)\s+maddedeki",
            rf"{article_no}\s*(inci|ıncı|uncu|üncü|nci|ncı|ncu|ncü)\s+maddede",
            rf"{article_no}\s*(inci|ıncı|uncu|üncü|nci|ncı|ncu|ncü)\s+maddenin",
        ]

    return 1.0 if any(re.search(pattern, text) for pattern in patterns) else 0.0


def extract_article_no_from_text(text: str) -> str:
    head = str(text)[:400].lower()
    patterns = [
        r"geçici\s+madde\s+(\d+)",
        r"gecici\s+madde\s+(\d+)",
        r"madde\s+([0-9]+)\s*[–\-:]",
        r"madde\s+([0-9]+)\b",
        r"^\s*([0-9]+)\.\s*madde",
    ]
    for pattern in patterns:
        match = re.search(pattern, head, flags=re.IGNORECASE | re.MULTILINE)
        if match:
            return match.group(1)
    return ""


def split_article_blocks(text: str) -> List[str]:
    text = normalize_whitespace(text)
    if not text:
        return []

    # Split on legal article headings while keeping the heading inside the block.
    pattern = re.compile(
        r"(?im)(?=^\s*(?:GEÇİCİ\s+MADDE|GECICI\s+MADDE|MADDE)\s+\d+\s*[–\-:]?)"
    )
    parts = [p.strip() for p in pattern.split(text) if p.strip()]

    # If the document does not look article-structured, fall back to character chunks.
    if len(parts) <= 1:
        return character_chunks(text, CHUNK_SIZE, CHUNK_OVERLAP)

    chunks: List[str] = []
    for part in parts:
        if len(part) <= CHUNK_SIZE * 1.5:
            chunks.append(part)
        else:
            chunks.extend(character_chunks(part, CHUNK_SIZE, CHUNK_OVERLAP))
    return chunks


def character_chunks(text: str, chunk_size: int, chunk_overlap: int) -> List[str]:
    text = normalize_whitespace(text)
    if len(text) <= chunk_size:
        return [text]

    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(text):
            break
        start = max(0, end - chunk_overlap)
    return chunks


def build_documents(uploaded_files, include_repo_docs: bool) -> List[Dict[str, str]]:
    documents: List[Dict[str, str]] = []

    if include_repo_docs and DEFAULT_DOCS_DIR.exists():
        for path in sorted(DEFAULT_DOCS_DIR.rglob("*")):
            if path.is_file() and not path.name.startswith(".") and path.suffix.lower() in [".txt", ".pdf", ".docx"]:
                text = read_file_path(path).strip()
                if text:
                    documents.append({
                        "doc_id": path.stem.replace(" ", "_"),
                        "file_name": path.name,
                        "source": infer_source_name(path.name, text),
                        "text": text,
                    })

    for uploaded in uploaded_files or []:
        data = uploaded.getvalue()
        text = read_file_bytes(uploaded.name, data).strip()
        if text:
            safe_stem = Path(uploaded.name).stem.replace(" ", "_")
            documents.append({
                "doc_id": f"uploaded_{safe_stem}",
                "file_name": uploaded.name,
                "source": infer_source_name(uploaded.name, text),
                "text": text,
            })

    if not documents:
        raise ValueError("Hiç doküman bulunamadı. Lütfen .txt, .pdf veya .docx hukuk dokümanı yükle.")

    return documents


def build_chunks_from_documents(documents: List[Dict[str, str]]) -> pd.DataFrame:
    rows = []
    for doc in documents:
        blocks = split_article_blocks(doc["text"])
        for idx, block in enumerate(blocks):
            if not block.strip():
                continue
            article_no = extract_article_no_from_text(block)
            rows.append({
                "chunk_id": f"{doc['doc_id']}_chunk_{idx:04d}",
                "doc_id": doc["doc_id"],
                "file_name": doc["file_name"],
                "source": doc["source"],
                "article_no": article_no,
                "chunk_text": block,
            })

    if not rows:
        raise ValueError("Dokümanlar yüklendi ama chunk oluşturulamadı.")

    return pd.DataFrame(rows)


# ----------------------------------
# Model loading
# ----------------------------------

@st.cache_resource(show_spinner=False)
def load_embedding_model() -> SentenceTransformer:
    return SentenceTransformer(EMBEDDING_MODEL_NAME)


@st.cache_resource(show_spinner=False)
def load_turkish_bge_reranker() -> CrossEncoder:
    return CrossEncoder(TURKISH_BGE_RERANKER_NAME, device="cpu", max_length=512)


@st.cache_resource(show_spinner=False)
def load_base_mistral(hf_token: str = ""):
    if torch is None or AutoTokenizer is None or AutoModelForCausalLM is None:
        raise ImportError(
            "Mistral generation için torch ve transformers gerekli: "
            "pip install torch transformers accelerate bitsandbytes"
        )

    token_arg = hf_token.strip() or os.environ.get("HF_TOKEN") or None

    tokenizer = AutoTokenizer.from_pretrained(
        BASE_MISTRAL_MODEL_NAME,
        use_fast=True,
        token=token_arg,
    )

    if tokenizer.pad_token is None:
        tokenizer.pad_token = tokenizer.eos_token
    tokenizer.padding_side = "right"
    tokenizer.truncation_side = "left"

    model_kwargs = {
        "low_cpu_mem_usage": True,
        "token": token_arg,
    }

    if torch.cuda.is_available():
        model_kwargs["device_map"] = "auto"
        model_kwargs["torch_dtype"] = torch.float16
        if BitsAndBytesConfig is not None:
            model_kwargs["quantization_config"] = BitsAndBytesConfig(
                load_in_4bit=True,
                bnb_4bit_quant_type="nf4",
                bnb_4bit_compute_dtype=torch.float16,
                bnb_4bit_use_double_quant=True,
            )
    else:
        # This may be slow and memory-heavy, but keeping it possible avoids hard failure
        # on CPU-only machines. For real demo, GPU/Colab is strongly recommended.
        model_kwargs["torch_dtype"] = torch.float32

    model = AutoModelForCausalLM.from_pretrained(BASE_MISTRAL_MODEL_NAME, **model_kwargs)
    model.eval()
    return tokenizer, model


# ----------------------------------
# Best-pipeline retrieval
# ----------------------------------

@st.cache_resource(show_spinner=False)
def build_retrieval_index(chunks_hash: str, chunks_records: Tuple[Tuple[str, str, str, str, str, str], ...]):
    chunks_df = pd.DataFrame(
        list(chunks_records),
        columns=["chunk_id", "doc_id", "file_name", "source", "article_no", "chunk_text"],
    )

    embedding_model = load_embedding_model()
    texts = chunks_df["chunk_text"].astype(str).tolist()
    embeddings = embedding_model.encode(
        texts,
        convert_to_numpy=True,
        show_progress_bar=False,
        normalize_embeddings=True,
    )

    tokenized_corpus = [simple_turkish_tokenize(text) for text in texts]
    bm25 = BM25Okapi(tokenized_corpus)

    return chunks_df, embedding_model, embeddings, bm25


def dataframe_to_hashable_records(chunks_df: pd.DataFrame) -> Tuple[Tuple[str, str, str, str, str, str], ...]:
    records = []
    for _, row in chunks_df.iterrows():
        records.append((
            str(row["chunk_id"]),
            str(row["doc_id"]),
            str(row["file_name"]),
            str(row["source"]),
            str(row.get("article_no", "")),
            str(row["chunk_text"]),
        ))
    return tuple(records)


def hash_records(records: Tuple[Tuple[str, str, str, str, str, str], ...]) -> str:
    digest = hashlib.sha256()
    for record in records:
        digest.update("|||".join(record).encode("utf-8", errors="ignore"))
    return digest.hexdigest()


def retrieve_candidates_article_source_aware(
    query: str,
    chunks_df: pd.DataFrame,
    embedding_model: SentenceTransformer,
    embeddings: np.ndarray,
    bm25: BM25Okapi,
    candidate_k: int = CANDIDATE_K,
    alpha: float = ALPHA,
    article_bonus_weight: float = ARTICLE_BONUS_WEIGHT,
) -> List[Dict[str, object]]:
    source_filter = detect_source_filter(query)
    candidate_df = chunks_df.copy().reset_index(drop=True)
    candidate_embeddings = embeddings
    used_source_filter = None

    if source_filter:
        mask = candidate_df.apply(lambda row: source_matches(source_filter, row), axis=1)
        filtered_df = candidate_df[mask].reset_index(drop=True)
        if len(filtered_df) >= MIN_FILTERED_ROWS:
            original_indices = candidate_df[mask].index.to_numpy()
            candidate_df = filtered_df
            candidate_embeddings = embeddings[original_indices]
            used_source_filter = source_filter

    candidate_texts = candidate_df["chunk_text"].astype(str).tolist()
    tokenized_candidate_corpus = [simple_turkish_tokenize(text) for text in candidate_texts]
    candidate_bm25 = BM25Okapi(tokenized_candidate_corpus)

    query_embedding = embedding_model.encode(
        [query],
        convert_to_numpy=True,
        normalize_embeddings=True,
    )

    dense_scores = cosine_similarity(query_embedding, candidate_embeddings)[0]
    bm25_scores = np.asarray(candidate_bm25.get_scores(simple_turkish_tokenize(query)), dtype=np.float32)

    dense_norm = min_max_normalize(dense_scores)
    bm25_norm = min_max_normalize(bm25_scores)
    hybrid_scores = alpha * dense_norm + (1.0 - alpha) * bm25_norm

    article_bonuses = np.asarray([
        article_match_bonus(query, candidate_df.iloc[i]["chunk_text"], candidate_df.iloc[i].get("article_no", ""))
        for i in range(len(candidate_df))
    ], dtype=np.float32)

    final_scores = hybrid_scores + article_bonus_weight * article_bonuses

    top_n = min(candidate_k, len(candidate_df))
    top_indices = np.argsort(final_scores)[::-1][:top_n]

    results = []
    for rank, idx in enumerate(top_indices, start=1):
        row = candidate_df.iloc[int(idx)]
        results.append({
            "rank": rank,
            "chunk_id": row["chunk_id"],
            "file_name": row["file_name"],
            "source": row["source"],
            "article_no": row.get("article_no", ""),
            "score": float(final_scores[int(idx)]),
            "hybrid_score": float(hybrid_scores[int(idx)]),
            "dense_score": float(dense_norm[int(idx)]),
            "bm25_score": float(bm25_norm[int(idx)]),
            "article_bonus": float(article_bonuses[int(idx)]),
            "source_filter": used_source_filter,
            "chunk_text": row["chunk_text"],
        })

    return results


def rerank_retrieved_chunks_turkish_bge(question: str, retrieved_results: List[Dict[str, object]]) -> List[Dict[str, object]]:
    if not retrieved_results:
        return []

    reranker = load_turkish_bge_reranker()
    pairs = [[question, item["chunk_text"]] for item in retrieved_results]
    scores = reranker.predict(pairs, batch_size=4, show_progress_bar=False)

    scored_results = []
    for item, score in zip(retrieved_results, scores):
        enriched = dict(item)
        enriched.update({
            "original_rank": item["rank"],
            "original_hybrid_score": item.get("hybrid_score", item.get("score", 0.0)),
            "rerank_score": float(score),
        })
        scored_results.append(enriched)

    scored_results = sorted(scored_results, key=lambda x: x["rerank_score"], reverse=True)

    for new_rank, item in enumerate(scored_results, start=1):
        item["rerank_rank"] = new_rank

    return scored_results


def retrieve_best_pipeline(
    question: str,
    chunks_df: pd.DataFrame,
    embedding_model: SentenceTransformer,
    embeddings: np.ndarray,
    bm25: BM25Okapi,
) -> List[Dict[str, object]]:
    candidates = retrieve_candidates_article_source_aware(
        query=question,
        chunks_df=chunks_df,
        embedding_model=embedding_model,
        embeddings=embeddings,
        bm25=bm25,
        candidate_k=CANDIDATE_K,
        alpha=ALPHA,
        article_bonus_weight=ARTICLE_BONUS_WEIGHT,
    )

    reranked_all = rerank_retrieved_chunks_turkish_bge(question, candidates)

    rerank_rank_map = {item["chunk_id"]: item["rerank_rank"] for item in reranked_all}
    rerank_score_map = {item["chunk_id"]: item["rerank_score"] for item in reranked_all}

    fused_results = []
    for item in candidates:
        original_rank = item["rank"]
        rerank_rank = rerank_rank_map.get(item["chunk_id"], CANDIDATE_K + 1)

        hybrid_rank_score = 1 / original_rank
        rerank_rank_score = 1 / rerank_rank
        fusion_score = HYBRID_WEIGHT * hybrid_rank_score + RERANK_WEIGHT * rerank_rank_score

        fused = dict(item)
        fused.update({
            "rank": None,
            "fusion_score": float(fusion_score),
            "original_rank": original_rank,
            "rerank_rank": rerank_rank,
            "rerank_score": float(rerank_score_map.get(item["chunk_id"], 0.0)),
        })
        fused_results.append(fused)

    fused_results = sorted(fused_results, key=lambda x: x["fusion_score"], reverse=True)

    final = []
    for rank, item in enumerate(fused_results[:FINAL_CONTEXT_K], start=1):
        item["rank"] = rank
        final.append(item)

    return final


# ----------------------------------
# Improved prompt + generation
# ----------------------------------

def build_improved_legal_rag_prompt(question: str, retrieved_contexts: List[str]) -> str:
    context_text = "\n\n".join([
        f"[Bağlam {i + 1}]\n{ctx}"
        for i, ctx in enumerate(retrieved_contexts)
    ])

    prompt = f"""
Sen Türk hukuk metinleri üzerinde çalışan dikkatli bir RAG soru-cevap asistanısın.

Görevin:
Sadece verilen bağlamları kullanarak soruya kısa, net ve hukuki olarak doğru cevap vermek.

Zorunlu kurallar:
1. Cevabı yalnızca verilen bağlamlara dayandır.
2. Bağlamda açıkça bulunmayan bilgiyi uydurma.
3. Soru bir süre, tarih, sayı veya madde soruyorsa sadece ilgili değeri ve kısa açıklamasını ver.
4. Soru "aykırı mıdır", "çelişir mi", "uygun mudur" gibi bir değerlendirme soruyorsa cevaba mutlaka "Evet" veya "Hayır" ile başla.
5. Cevapta "Bağlam", "Context", "verilen metne göre" gibi ifadeler kullanma.
6. Cevap en fazla iki kısa cümle olmalı.
7. Alternatif cevap, yeni soru, örnek soru, başlık veya açıklama bölümü üretme.
8. Cevabı verdikten sonra dur.
9. Eğer bağlamda cevap yoksa sadece şunu yaz: "Verilen bağlamda bu sorunun cevabı bulunamamaktadır."

Bağlamlar:
{context_text}

Soru:
{question}

Kısa ve doğrudan cevap:
"""

    # Mistral Instruct format.
    return f"<s>[INST] {prompt.strip()} [/INST]"


def clean_generated_answer(text: str) -> str:
    text = str(text).strip()

    start_markers = [
        "Kısa ve doğrudan cevap:",
        "Kısa cevap:",
        "Answer:",
        "Cevap:",
    ]
    for marker in start_markers:
        if marker in text:
            text = text.split(marker)[-1].strip()

    stop_markers = [
        "\nBaşka bir şekilde:",
        "\nCevap:",
        "\nSoru:",
        "\nQuestion:",
        "\nAlternatif",
        "\nAçıklama:",
        "\nDetaylı cevap:",
        "\nÖrnek:",
        "</s>",
    ]
    for marker in stop_markers:
        if marker in text:
            text = text.split(marker)[0].strip()

    unwanted = ["Context 1", "Context 2", "Context 3", "Bağlam 1", "Bağlam 2", "Bağlam 3"]
    for u in unwanted:
        text = text.replace(u, "")

    text = re.sub(r"\(\s*\)", "", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def postprocess_answer_by_question_type(question: str, answer: str) -> str:
    q = str(question).lower()
    answer = str(answer).strip()

    short_answer_triggers = ["kaç", "ne zaman", "kime aittir", "en fazla kaç"]
    if any(trigger in q for trigger in short_answer_triggers):
        sentences = re.split(r"(?<=[.!?])\s+", answer)
        if sentences:
            return sentences[0].strip()

    return answer


def generate_with_base_mistral(prompt: str, hf_token: str = "", max_new_tokens: int = 120) -> str:
    tokenizer, model = load_base_mistral(hf_token)

    inputs = tokenizer(
        prompt,
        return_tensors="pt",
        truncation=True,
        max_length=2048,
    )

    if torch is not None and torch.cuda.is_available():
        inputs = inputs.to(model.device)

    input_length = inputs["input_ids"].shape[1]

    with torch.no_grad():
        outputs = model.generate(
            **inputs,
            max_new_tokens=max_new_tokens,
            do_sample=False,
            temperature=0.0,
            pad_token_id=tokenizer.eos_token_id,
            eos_token_id=tokenizer.eos_token_id,
        )

    new_tokens = outputs[0][input_length:]
    answer = tokenizer.decode(new_tokens, skip_special_tokens=True)
    return answer.strip()


# ----------------------------------
# Streamlit UI
# ----------------------------------

def main() -> None:
    st.set_page_config(
        page_title="Turkish Legal RAG Demo",
        page_icon="⚖️",
        layout="wide",
    )

    st.markdown(
        """
        <style>
        .block-container {padding-top: 2rem; max-width: 1100px;}
        .small-muted {color: #666; font-size: 0.9rem;}
        div[data-testid="stExpander"] {border-radius: 12px;}
        </style>
        """,
        unsafe_allow_html=True,
    )

    st.title("⚖️ Turkish Legal RAG")
    st.caption("Pipeline: Source-aware + article-aware retrieval, Turkish BGE reranker, improved legal prompt, base Mistral generator")

    with st.sidebar:
        st.subheader("System Status")
        if torch is not None and torch.cuda.is_available():
            st.success(f"GPU available: {torch.cuda.get_device_name(0)}")
        else:
            st.warning("GPU görünmüyor. Base Mistral generation CPU'da çok yavaş olabilir veya belleğe sığmayabilir.")

        hf_token = st.text_input(
            "Hugging Face token (optional)",
            type="password",
            help="Mistral modeli indirilirken gerekirse kullanılır. HF_TOKEN environment variable varsa boş bırakabilirsin.",
        )

        include_repo_docs = st.checkbox(
            "Include existing final_system documents",
            value=True,
            help="Kapalı yaparsan sadece yüklenen dosyalara göre cevap verir.",
        )

        st.markdown("---")
        st.caption("UI sade tutuldu: tek soru → tek generated answer. Kaynaklar sadece kontrol için aşağıda açılır.")

    st.subheader("1. Upload Legal Document(s)")
    uploaded_files = st.file_uploader(
        "Hocanın kendi hukuk dokümanlarını buraya yükleyebilirsin (.txt, .pdf, .docx).",
        type=SUPPORTED_UPLOAD_TYPES,
        accept_multiple_files=True,
    )

    try:
        documents = build_documents(uploaded_files, include_repo_docs=include_repo_docs)
        chunks_df = build_chunks_from_documents(documents)

        st.success(f"{len(documents)} document(s) loaded, {len(chunks_df)} chunk(s) prepared.")
    except Exception as exc:
        st.error(str(exc))
        return

    st.subheader("2. Ask a Question")
    question = st.text_area(
        "Question",
        value="Anayasa 10. madde neyi düzenler?",
        height=110,
        label_visibility="collapsed",
        placeholder="Örn: Anayasa 10. madde neyi düzenler?",
    )

    generate_clicked = st.button("Generate Answer", type="primary", use_container_width=True)

    if generate_clicked:
        if not question.strip():
            st.warning("Lütfen bir soru yaz.")
            return

        try:
            records = dataframe_to_hashable_records(chunks_df)
            records_hash = hash_records(records)

            with st.spinner("Embedding index is being prepared..."):
                indexed_chunks_df, embedding_model, embeddings, bm25 = build_retrieval_index(records_hash, records)

            with st.spinner("Retrieving and reranking legal contexts with Turkish BGE..."):
                retrieved = retrieve_best_pipeline(question, indexed_chunks_df, embedding_model, embeddings, bm25)

            if not retrieved:
                st.error("Soru için bağlam getirilemedi.")
                return

            contexts = [item["chunk_text"] for item in retrieved]
            prompt = build_improved_legal_rag_prompt(question, contexts)

            with st.spinner("Base Mistral is generating the answer..."):
                raw_answer = generate_with_base_mistral(prompt, hf_token=hf_token)
                answer = clean_generated_answer(raw_answer)
                answer = postprocess_answer_by_question_type(question, answer)

            st.subheader("Generated Answer")
            st.write(answer if answer else "Cevap üretilemedi.")

            with st.expander("Retrieved Legal Sources", expanded=False):
                for item in retrieved:
                    st.markdown(f"**Rank {item['rank']} — {item['source']} / {item['file_name']}**")
                    if item.get("article_no"):
                        st.caption(f"Article: {item['article_no']}")
                    st.caption(
                        f"fusion={item.get('fusion_score', 0):.4f} | "
                        f"rerank={item.get('rerank_score', 0):.4f} | "
                        f"article_bonus={item.get('article_bonus', 0):.2f} | "
                        f"source_filter={item.get('source_filter') or '-'}"
                    )
                    st.write(item["chunk_text"])
                    st.markdown("---")

        except Exception as exc:
            st.error("Answer generation failed.")
            st.exception(exc)
            st.info(
                "Bu pipeline gerçek base Mistral modelini yükler. Çalışması için genelde GPU/Colab ve gerekirse Hugging Face access gerekir. "
                "Retrieval + reranker tarafı çalışsa bile Mistral yüklenemezse cevap üretimi tamamlanamaz."
            )


if __name__ == "__main__":
    main()
